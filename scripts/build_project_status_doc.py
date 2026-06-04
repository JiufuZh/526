from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Group8_Defect_LLM_Project_Status_Report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if "\n" not in text and len(text) < 20 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr_cells[i], header, bold=True, color="FFFFFF")
        set_cell_shading(hdr_cells[i], "1F4D78")
        if widths:
            hdr_cells[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(31, 77, 120) if level <= 2 else RGBColor(67, 67, 67)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(10.5)
        rest = text[len(bold_prefix):]
        r2 = p.add_run(rest)
        r2.font.name = "Arial"
        r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.5)
for style_name, size, color in [
    ("Heading 1", 16, "1F4D78"),
    ("Heading 2", 13, "2E74B5"),
    ("Heading 3", 12, "434343"),
]:
    style = styles[style_name]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Group 8 Defect LLM Project\n当前进展、结果与运行逻辑汇总")
r.font.name = "Arial"
r.font.size = Pt(20)
r.bold = True
r.font.color.rgb = RGBColor(11, 37, 69)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run("更新日期：2026-06-01 | 项目路径：/gpfs/projects/imt526a/group8/final")
r.font.name = "Arial"
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(85, 85, 85)

add_heading(doc, "1. 项目目标", 1)
add_body(
    doc,
    "这个项目的核心目标是：在 CodeXGLUE/Devign 风格的软件缺陷检测任务上，比较不同 LLM/模型方案对“代码是否有缺陷”的二分类效果，并形成可以写进报告的实验证据。",
)
add_body(
    doc,
    "目前主线不是只看 accuracy，而是同时看 macro-F1 和 defective-F1，因为数据集类别并不完全均衡，而且模型如果只预测 non_defective，也可能得到看似不低的 accuracy，但对真正有缺陷的代码没有识别能力。",
)

add_heading(doc, "2. 已经完成的内容", 1)
for item in [
    "完成了项目环境和运行路径整理，Tillicum 项目目录为 /gpfs/projects/imt526a/group8/final。",
    "完成了 LoRA fine-tuning 主实验，产物为 outputs/qwen25_7b_lora_bf16_512/final_adapter。",
    "完成了 fine-tuned LoRA 在 validation 和 test split 上的评估，并生成 metrics/predictions 文件。",
    "完成了 zero-shot baseline，用于衡量不训练时模型的原始能力。",
    "完成了 4-shot few-shot baseline，但结果显示模型退化为几乎全预测 non_defective，因此不能作为有力提升证据。",
    "完成了 validation/test 的错误分析文件，包含误分类样本和错误类型摘要。",
    "已经启动下一步 GraphCodeBERT encoder baseline，用来补强和传统代码模型的对比。",
]:
    add_bullet(doc, item)

add_heading(doc, "3. 当前主要结果", 1)
add_body(doc, "下面这张表是目前最关键的可报告结果。数字为已完成实验的当前记录，GraphCodeBERT 仍在运行/排队监控中。")
add_table(
    doc,
    ["方法", "Validation", "Test", "结论"],
    [
        [
            "Zero-shot Qwen",
            "Acc 0.5201\nMacro-F1 0.5086\nDefective-F1 0.4332",
            "Acc 0.5264\nMacro-F1 0.5180\nDefective-F1 0.4545",
            "无训练基线；能做出一些判断，但缺陷类识别仍弱。",
        ],
        [
            "4-shot Qwen",
            "Acc 0.5655\nMacro-F1 0.3612\nDefective-F1 0.0000",
            "Acc 0.5406\nMacro-F1 0.3509\nDefective-F1 0.0000",
            "accuracy 看起来不差，但实际全偏向 non_defective，不适合作为主要结果。",
        ],
        [
            "LoRA fine-tuned Qwen",
            "Acc 0.5556\nMacro-F1 0.5509\nDefective-F1 0.5045",
            "Acc 0.5523\nMacro-F1 0.5507\nDefective-F1 0.5236",
            "当前最稳的 LLM 结果；相比 zero-shot，test macro-F1 提升约 +0.0327。",
        ],
        [
            "GraphCodeBERT baseline",
            "运行/排队中",
            "待生成",
            "用于回答“传统代码 encoder 是否比 LLM fine-tuning 更强”。",
        ],
    ],
    widths=[Inches(1.45), Inches(1.45), Inches(1.45), Inches(2.65)],
)

add_heading(doc, "4. Fine-tuning 前后对比", 1)
add_body(
    doc,
    "Fine-tuning 的提升不是特别巨大，但它是稳定且方向正确的：zero-shot test macro-F1 约为 0.5180，LoRA fine-tuned test macro-F1 约为 0.5507，提升约 +0.0327；defective-F1 从 0.4545 提升到 0.5236，说明模型对缺陷类的识别能力确实变好。",
)
add_body(
    doc,
    "相比之下，4-shot 的 accuracy 有 0.5406，但 defective recall 和 defective-F1 都是 0，说明它几乎没有识别任何 defective 样本。因此报告里应该强调：accuracy 不能单独说明模型好坏，macro-F1/defective-F1 才是更有说服力的指标。",
)

add_heading(doc, "5. 现在正在做什么", 1)
add_body(
    doc,
    "当前正在跑 GraphCodeBERT encoder baseline。这个实验是为了增加一个非生成式、代码预训练 encoder 的对照组，避免最终报告只有 prompt/LoRA LLM 之间的比较。",
)
add_bullet(doc, "Slurm job id：131579。")
add_bullet(doc, "脚本路径：outputs/run_graphcodebert_1h.slurm。")
add_bullet(doc, "提交记录：outputs/graphcodebert_sbatch.txt。")
add_bullet(doc, "预期输出：outputs/graphcodebert_defect_baseline/validation_metrics.json。")
add_bullet(doc, "日志位置：outputs/slurm/graphcodebert-*.out 和 outputs/slurm/graphcodebert-*.err。")
add_body(
    doc,
    "如果这个任务还在 pending，最常见原因是 Tillicum 账户同时可用的 GPU job 数有限，当前 OOD GPU session 占着额度时，Slurm job 会等待。等额度释放后它会继续排队/启动。",
)

add_heading(doc, "6. 各步骤之间的逻辑", 1)
for step in [
    "先跑 zero-shot：建立“不训练时模型能做到什么”的最低成本基线。",
    "再跑 LoRA fine-tuning：验证 proposal 中的核心假设，即通过小参数量微调让模型更适应缺陷检测任务。",
    "再做 validation/test 评估：validation 用来观察模型选择是否合理，test 用来报告最终泛化结果。",
    "再做 error analysis：把错误样本整理出来，方便报告解释模型为什么错、错在哪里。",
    "再跑 few-shot：测试 prompt 中给少量样例是否能替代训练；当前结果说明 few-shot 在这个任务上不稳定。",
    "最后补 GraphCodeBERT：加入传统代码 encoder baseline，让最终结论更有说服力，也给提高 F1 留一个实际方向。",
]:
    add_number(doc, step)

add_heading(doc, "7. 结果文件在哪里看", 1)
add_table(
    doc,
    ["内容", "路径"],
    [
        ["LoRA adapter", "outputs/qwen25_7b_lora_bf16_512/final_adapter"],
        ["LoRA validation metrics", "outputs/qwen25_7b_lora_bf16_512/validation_metrics.json"],
        ["LoRA test metrics", "outputs/qwen25_7b_lora_bf16_512/test_metrics.json"],
        ["LoRA predictions", "outputs/qwen25_7b_lora_bf16_512/validation_predictions.csv / test_predictions.csv"],
        ["Zero-shot metrics", "outputs/zero_fewshot_0_shot_validation_metrics.json / outputs/zero_fewshot_0_shot_test_metrics.json"],
        ["4-shot metrics", "outputs/zero_fewshot_4_shot_validation_metrics.json / outputs/zero_fewshot_4_shot_test_metrics.json"],
        ["Report-ready summary", "outputs/report_ready_metrics.csv / outputs/report_ready_metrics.md"],
        ["Error analysis", "outputs/qwen25_7b_lora_bf16_512/error_analysis_validation/ 和 error_analysis_test/"],
        ["GraphCodeBERT logs", "outputs/slurm/graphcodebert-*.out / graphcodebert-*.err"],
    ],
    widths=[Inches(2.1), Inches(4.8)],
)

add_heading(doc, "8. 目前是否完成 proposal", 1)
add_body(
    doc,
    "从 proposal 的主线来看，核心 fine-tuning 实验、zero-shot 对比、test evaluation 和 error analysis 已经完成。现在额外补 GraphCodeBERT baseline，是为了让最终报告更完整、更有说服力：如果 GraphCodeBERT 更强，可以作为 strong baseline；如果 LoRA 接近或超过它，也能强化 fine-tuning 的价值。",
)
add_body(
    doc,
    "因此当前状态可以概括为：主实验已完成，报告所需核心结果已具备；正在补充更强 baseline 和可解释对比，以提升最终呈现质量。",
)

add_heading(doc, "9. 下一步建议", 1)
for item in [
    "继续监控 GraphCodeBERT，拿到 validation/test 指标后加入总表。",
    "报告里优先展示 macro-F1 和 defective-F1，accuracy 作为辅助指标。",
    "把 4-shot 作为负面/诊断结果：说明 naive few-shot prompt 对该任务不可靠。",
    "如果时间允许，再尝试阈值调整、class weight、或更长/更稳的 encoder baseline，以争取进一步提升 defective-F1。",
]:
    add_bullet(doc, item)

footer = doc.sections[0].footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("Group 8 Defect LLM Project Status Report")
fr.font.name = "Arial"
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(85, 85, 85)

doc.save(OUT)
print(OUT)
