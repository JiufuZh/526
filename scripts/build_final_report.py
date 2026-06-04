from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"
DOCS = ROOT / "docs"
PDF_OUT = DOCS / "Group8_Final_Report.pdf"
DOCX_OUT = DOCS / "Group8_Final_Report.docx"


TEAM_MEMBERS = "[fill names in alphabetical order by last name]"
CODE_LINK = "https://github.com/JiufuZh/526"


def load_metric(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def metric_row(name: str, path: str) -> dict:
    m = load_metric(RESULTS / path)
    return {
        "method": name,
        "accuracy": m["accuracy"],
        "macro_f1": m["macro_f1"],
        "defective_f1": m["defective_f1"],
        "defective_recall": m["defective_recall"],
        "tn": m["tn"],
        "fp": m["fp"],
        "fn": m["fn"],
        "tp": m["tp"],
    }


TEST_ROWS = [
    metric_row("GraphCodeBERT", "graphcodebert_test_metrics.json"),
    metric_row("TF-IDF Logistic Regression", "tfidf_logreg_test_metrics.json"),
    metric_row("TF-IDF Linear SVM", "tfidf_linear_svm_test_metrics.json"),
    metric_row("Qwen LoRA fine-tuned", "lora_test_metrics.json"),
    metric_row("Qwen zero-shot", "zero_shot_test_metrics.json"),
    metric_row("Qwen 4-shot", "four_shot_test_metrics.json"),
    metric_row("Majority baseline", "majority_test_metrics.json"),
]

VALIDATION_ROWS = [
    metric_row("GraphCodeBERT", "graphcodebert_validation_metrics.json"),
    metric_row("Qwen LoRA fine-tuned", "lora_validation_metrics.json"),
    metric_row("Qwen zero-shot", "zero_shot_validation_metrics.json"),
    metric_row("Qwen 4-shot", "four_shot_validation_metrics.json"),
]


def fmt(x: float) -> str:
    return f"{x:.4f}"


def para(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(text.replace("&", "&amp;"), style)


def bullets(items: list[str], style: ParagraphStyle) -> ListFlowable:
    return ListFlowable(
        [ListItem(para(item, style), leftIndent=12) for item in items],
        bulletType="bullet",
        start="circle",
        leftIndent=18,
        bulletFontSize=7,
    )


def make_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "Title",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=18,
            leading=22,
            alignment=TA_CENTER,
            spaceAfter=8,
            textColor=colors.HexColor("#0B2545"),
        ),
        "subtitle": ParagraphStyle(
            "Subtitle",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=13,
            alignment=TA_CENTER,
            spaceAfter=12,
            textColor=colors.HexColor("#333333"),
        ),
        "h1": ParagraphStyle(
            "Heading1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            spaceBefore=10,
            spaceAfter=5,
            textColor=colors.HexColor("#2E74B5"),
        ),
        "h2": ParagraphStyle(
            "Heading2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11,
            leading=14,
            spaceBefore=7,
            spaceAfter=4,
            textColor=colors.HexColor("#1F4D78"),
        ),
        "body": ParagraphStyle(
            "Body",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9.6,
            leading=12.2,
            spaceAfter=5,
            alignment=TA_LEFT,
        ),
        "small": ParagraphStyle(
            "Small",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=8.3,
            leading=10.2,
            spaceAfter=3,
            alignment=TA_LEFT,
        ),
        "caption": ParagraphStyle(
            "Caption",
            parent=base["BodyText"],
            fontName="Helvetica-Oblique",
            fontSize=8,
            leading=10,
            spaceBefore=2,
            spaceAfter=5,
            textColor=colors.HexColor("#555555"),
        ),
    }


def styled_table(data: list[list[str]], widths: list[float], header_rows: int = 1) -> Table:
    table = Table(data, colWidths=widths, repeatRows=header_rows, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 7.6),
                ("LEADING", (0, 0), (-1, -1), 9.2),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0B2545")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C8D0DA")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (1, 1), (-1, -1), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#666666"))
    canvas.drawString(inch, 0.55 * inch, "Group 8 Code Defect Detection Final Report")
    canvas.drawRightString(7.5 * inch, 0.55 * inch, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf() -> None:
    DOCS.mkdir(exist_ok=True)
    styles = make_styles()
    story = []

    story.append(para("Code Defect Detection with LLM Prompting, LoRA Fine-Tuning, and GraphCodeBERT", styles["title"]))
    story.append(
        para(
            f"Group 8 | Team members: {TEAM_MEMBERS}<br/>Course final project report | Code: {CODE_LINK}",
            styles["subtitle"],
        )
    )
    story.append(para("Abstract", styles["h1"]))
    story.append(
        para(
            "This project studies function-level code defect detection as a binary classification task: given a C/C++ function, predict whether it is non-defective or defective. The main research question is whether prompt-only large language models are sufficient for this task, or whether supervised adaptation and code-specific pretraining are necessary. We evaluate a ladder of methods: Qwen2.5-7B-Instruct zero-shot prompting, Qwen 4-shot prompting, Qwen LoRA fine-tuning, majority and TF-IDF CPU baselines, and a GraphCodeBERT encoder classifier. On the held-out test split, GraphCodeBERT achieves the best balanced performance with 0.6589 accuracy and 0.6517 Macro-F1. TF-IDF Logistic Regression is also strong, with 0.6218 Macro-F1 and the highest Defective-F1 of 0.6088. Prompt-only Qwen is weaker, and 4-shot prompting collapses to all non-defective predictions. These results support the conclusion that code-aware supervised models are more reliable than prompt-only LLM use for defect detection.",
            styles["body"],
        )
    )

    story.append(para("1. Problem, Motivation, and Novelty", styles["h1"]))
    story.append(
        para(
            "Automated defect detection aims to identify functions that may contain bugs or vulnerabilities before deployment. The motivation is practical: manual code review is costly, defects can be subtle, and false negatives can leave risky code undetected. The project is novel for the course setting because it does not simply fine-tune one model. Instead, it compares a full adaptation ladder from no-training LLM prompting, to in-context prompting, to parameter-efficient LLM fine-tuning, to traditional lexical baselines, to a code-specific transformer encoder. This design lets us analyze not only which model performs best, but what kind of adaptation actually helps.",
            styles["body"],
        )
    )

    story.append(para("2. Related Work and Differentiation", styles["h1"]))
    story.append(
        para(
            "The dataset and task follow the CodeXGLUE defect detection benchmark, which is based on the Devign-style function-level defect setting [1, 2]. GraphCodeBERT is a code-pretrained transformer that incorporates code structure signals and is a natural supervised encoder baseline for this task [3]. Qwen2.5-7B-Instruct represents a modern instruction-tuned LLM family that can be evaluated through prompts [4]. LoRA provides a lightweight way to adapt large models without full fine-tuning [5]. Our differentiation is the controlled comparison across these families under the same train/validation/test splits and the same evaluation metrics. The project therefore contrasts prompt-only behavior, parameter-efficient LLM adaptation, lexical supervised baselines, and code-specific supervised pretraining.",
            styles["body"],
        )
    )

    story.append(para("3. Dataset and Methods", styles["h1"]))
    story.append(
        para(
            "We use the CodeXGLUE defect detection benchmark. Labels are binary: 0 means non-defective and 1 means defective. The pipeline verifies splits, prepares cached datasets, runs prompt-based LLM inference, trains supervised baselines, evaluates validation and test predictions, and saves metrics plus confusion matrices. The implementation is hosted in the project GitHub repository and the main Tillicum run directory was /gpfs/projects/imt526a/group8/final.",
            styles["body"],
        )
    )
    method_data = [
        ["Method", "Purpose", "Adaptation"],
        ["Majority", "Sanity-check baseline", "Most frequent class"],
        ["TF-IDF + SVM/LogReg", "Traditional supervised lexical baseline", "CPU classifier"],
        ["Qwen zero-shot", "Raw instruction-following LLM reasoning", "Prompt only"],
        ["Qwen 4-shot", "In-context examples without training", "Prompt examples"],
        ["Qwen LoRA", "Supervised LLM task adaptation", "LoRA adapter"],
        ["GraphCodeBERT", "Code-specific supervised encoder", "Classifier head training"],
    ]
    story.append(styled_table(method_data, [1.45 * inch, 3.25 * inch, 1.65 * inch]))

    story.append(
        KeepTogether(
            [
                para("4. Evaluation Measures", styles["h1"]),
                para(
                    "The evaluation uses accuracy, Macro-F1, Defective-F1, defective recall, and confusion matrices. Accuracy is useful but insufficient because a model can obtain moderate accuracy by predicting the majority class while missing all defective functions. Macro-F1 is more appropriate because it gives balanced weight to both classes. Defective-F1 and defective recall are especially important because the high-risk error is failing to catch defective code. The confusion matrix reports true negatives, false positives, false negatives, and true positives so the report can explain how each model behaves.",
                    styles["body"],
                ),
            ]
        )
    )

    story.append(para("5. Results", styles["h1"]))
    story.append(para("Table 1 summarizes the held-out test results, sorted by Macro-F1.", styles["caption"]))
    table_rows = [["Model", "Acc.", "Macro-F1", "Def.-F1", "Def. recall", "TN", "FP", "FN", "TP"]]
    for r in TEST_ROWS:
        table_rows.append(
            [
                r["method"],
                fmt(r["accuracy"]),
                fmt(r["macro_f1"]),
                fmt(r["defective_f1"]),
                fmt(r["defective_recall"]),
                str(r["tn"]),
                str(r["fp"]),
                str(r["fn"]),
                str(r["tp"]),
            ]
        )
    story.append(styled_table(table_rows, [1.55 * inch, 0.58 * inch, 0.68 * inch, 0.65 * inch, 0.72 * inch, 0.45 * inch, 0.45 * inch, 0.45 * inch, 0.45 * inch]))
    story.append(Spacer(1, 4))
    story.append(
        para(
            "GraphCodeBERT is the strongest balanced model on test, reaching 0.6517 Macro-F1. TF-IDF Logistic Regression is the strongest traditional baseline and is highly competitive, with 0.6088 Defective-F1. Qwen LoRA improves over zero-shot prompting on Macro-F1 and Defective-F1, but it does not reach the supervised lexical or code-specific baselines. The 4-shot model and majority baseline both predict every test sample as non-defective, giving Defective-F1 of 0.",
            styles["body"],
        )
    )

    fig_path = RESULTS / "figures" / "confusion_matrices_overview.png"
    if fig_path.exists():
        story.append(KeepTogether([Image(str(fig_path), width=6.2 * inch, height=3.3 * inch), para("Figure 1. Test/validation confusion matrices show class collapse for 4-shot and majority baselines, and more balanced behavior for supervised models.", styles["caption"])]))

    story.append(para("6. Analysis of Outcomes", styles["h1"]))
    story.append(
        bullets(
            [
                "Prompt-only LLMs are not reliable enough for this task. Zero-shot Qwen reaches 0.5180 test Macro-F1, while 4-shot prompting collapses to non-defective predictions.",
                "LoRA fine-tuning helps the LLM: test Macro-F1 increases from 0.5180 to 0.5507, and Defective-F1 increases from 0.4545 to 0.5236.",
                "Traditional supervised baselines are strong. TF-IDF Logistic Regression reaches 0.6218 Macro-F1 and 0.6088 Defective-F1, showing that lexical signals in the function text are meaningful.",
                "GraphCodeBERT provides the best balanced performance, suggesting that code-specific pretraining is useful when paired with supervised labels.",
            ],
            styles["body"],
        )
    )
    story.append(
        para(
            "The error-analysis summaries also suggest that false positives tend to occur on longer and more complex functions than false negatives. On the test split, false positives average 160.4 lines and 15.4 branch points, compared with 52.1 lines and 3.25 branch points for false negatives. This supports the interpretation that complexity and pointer-heavy code can trigger conservative defect predictions.",
            styles["body"],
        )
    )

    story.append(para("7. Demo Plan and Reproducibility", styles["h1"]))
    story.append(
        para(
            "The 10-minute demo should use the prepared notebook and repository artifacts rather than retraining live. A professional flow is: introduce the problem and research question, show the method ladder, open the report-ready metrics table, show the confusion matrix figure, briefly show the key scripts/config files, and close with the main takeaway. The demo notebook reads completed outputs, so it is stable for screen sharing. Code, scripts, configs, result JSON files, and generated figures are available at the GitHub link above.",
            styles["body"],
        )
    )

    story.append(para("8. Limitations and Future Work", styles["h1"]))
    story.append(
        bullets(
            [
                "The best model reaches about 0.65 Macro-F1, so the task is not solved.",
                "The LoRA run is parameter-efficient but may need threshold tuning, longer context, or more training budget.",
                "Few-shot prompting may improve with better example selection, but the current naive 4-shot setup is not reliable.",
                "Future work should test larger code encoders, calibration for defective recall, and deeper error analysis by defect type.",
            ],
            styles["body"],
        )
    )

    story.append(para("References", styles["h1"]))
    refs = [
        "[1] CodeXGLUE benchmark. https://github.com/microsoft/CodeXGLUE",
        "[2] Devign: Effective Vulnerability Identification by Learning Comprehensive Program Semantics via Graph Neural Networks. https://arxiv.org/abs/1909.03496",
        "[3] GraphCodeBERT: Pre-training Code Representations with Data Flow. https://arxiv.org/abs/2009.08366 and https://huggingface.co/microsoft/graphcodebert-base",
        "[4] Qwen model family. https://huggingface.co/Qwen",
        "[5] LoRA: Low-Rank Adaptation of Large Language Models. https://arxiv.org/abs/2106.09685",
        f"[6] Project repository and generated code/results. {CODE_LINK}",
    ]
    for ref in refs:
        story.append(para(ref, styles["small"]))

    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=letter,
        rightMargin=0.7 * inch,
        leftMargin=0.7 * inch,
        topMargin=0.65 * inch,
        bottomMargin=0.75 * inch,
        title="Group 8 Final Report",
        author="Group 8",
    )
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def add_docx_table(doc: Document, data: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(data[0]))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, value in enumerate(data[0]):
        hdr[i].text = value
    for row in data[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value


def build_docx() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for name, size in [("Heading 1", 15), ("Heading 2", 12)]:
        styles[name].font.name = "Calibri"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor(46, 116, 181)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Code Defect Detection with LLM Prompting, LoRA Fine-Tuning, and GraphCodeBERT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = doc.add_paragraph(f"Group 8 | Team members: {TEAM_MEMBERS}\nCourse final project report | Code: {CODE_LINK}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sections = [
        ("Abstract", "This project studies function-level code defect detection as a binary classification task. We compare prompt-only Qwen, Qwen LoRA fine-tuning, traditional CPU baselines, and GraphCodeBERT. GraphCodeBERT achieves the best balanced test performance with 0.6589 accuracy and 0.6517 Macro-F1."),
        ("Problem, Motivation, and Novelty", "The project asks whether prompt-only LLMs are enough for code defect detection or whether supervised adaptation and code-specific pretraining are necessary. The novelty is the full comparison ladder from no-training prompting to code-specific supervised modeling."),
        ("Related Work", "The work uses CodeXGLUE/Devign-style defect detection, Qwen2.5-7B-Instruct prompting, LoRA parameter-efficient fine-tuning, TF-IDF classifiers, and GraphCodeBERT."),
        ("Dataset and Methods", "The dataset is CodeXGLUE defect detection with binary labels: non-defective and defective. The pipeline prepares data, trains/evaluates models, and saves metrics plus confusion matrices."),
        ("Evaluation Measures", "Accuracy, Macro-F1, Defective-F1, defective recall, and confusion matrices are reported. Macro-F1 and Defective-F1 are emphasized because accuracy can hide class collapse."),
    ]
    for heading, text in sections:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(text)

    doc.add_heading("Results", level=1)
    result_data = [["Model", "Acc.", "Macro-F1", "Def.-F1", "Def. recall"]]
    for r in TEST_ROWS:
        result_data.append([r["method"], fmt(r["accuracy"]), fmt(r["macro_f1"]), fmt(r["defective_f1"]), fmt(r["defective_recall"])])
    add_docx_table(doc, result_data)

    doc.add_heading("Analysis and Future Work", level=1)
    doc.add_paragraph("GraphCodeBERT is the best balanced model. TF-IDF Logistic Regression is surprisingly competitive and has the highest Defective-F1. LoRA improves over zero-shot prompting, but prompt-only 4-shot collapses to non-defective predictions.")
    doc.add_paragraph("Future work should test threshold tuning, better few-shot example selection, larger code encoders, and richer error analysis.")

    doc.add_heading("References", level=1)
    for ref in [
        "CodeXGLUE benchmark: https://github.com/microsoft/CodeXGLUE",
        "Devign paper: https://arxiv.org/abs/1909.03496",
        "GraphCodeBERT paper/model: https://arxiv.org/abs/2009.08366 and https://huggingface.co/microsoft/graphcodebert-base",
        "Qwen model family: https://huggingface.co/Qwen",
        "LoRA paper: https://arxiv.org/abs/2106.09685",
        f"Project code: {CODE_LINK}",
    ]:
        doc.add_paragraph(ref, style="List Bullet")
    doc.save(DOCX_OUT)


def main() -> None:
    build_pdf()
    build_docx()
    print(PDF_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
