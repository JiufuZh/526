from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Group8_PPT_Generation_Brief.docx"
FIG = ROOT / "results" / "figures" / "confusion_matrices_overview.png"


BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2F3"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(row.cells[idx])


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = BLUE
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.35])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.color.rgb = BLUE
    p.add_run(text)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_GRAY)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            set_cell_border(cells[i])
    doc.add_paragraph()
    return table


def add_prompt_box(doc, prompt):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.35])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7FBFF")
    p = cell.paragraphs[0]
    for line in prompt.strip().splitlines():
        if p.text:
            p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    doc.add_paragraph()


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size in [("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11.5)]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Group 8 Defect Detection Project\nPPT Generation Brief")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Use this document as the source brief for GPT to generate a 10-minute final demo deck.").italic = True
    doc.add_paragraph()

    add_note(
        doc,
        "Fill before final PPT",
        "Add all team member names in alphabetical order by last name on the title slide and report. Current placeholder: Group 8 members: [fill names].",
    )

    add_heading(doc, "1. One-Paragraph Project Summary", 1)
    doc.add_paragraph(
        "This project evaluates whether large language models and code-specific encoders can detect defective C/C++ functions from the CodeXGLUE defect detection benchmark. "
        "We compare prompt-only Qwen, LoRA fine-tuned Qwen, traditional TF-IDF CPU baselines, and a supervised GraphCodeBERT encoder baseline. "
        "The main finding is that prompt-only LLM use is not reliable for this task, TF-IDF baselines are surprisingly competitive, and GraphCodeBERT achieves the strongest balanced test Macro-F1."
    )

    add_heading(doc, "2. Rubric-Aligned Goals", 1)
    add_bullets(
        doc,
        [
            "Demo: use the full 10 minutes with a planned flow, live screen share, and all group members present.",
            "Professionalism: show a clean GitHub repo, reproducible outputs, clear figures, and a short scripted demo path.",
            "Report: explain novelty/motivation, related work, evaluation measures, results quality, and presentation polish.",
            "Submission: one PDF report, no more than 8 pages, plus code link: https://github.com/JiufuZh/526.",
        ],
    )

    add_heading(doc, "3. Core Storyline for Slides", 1)
    add_numbered(
        doc,
        [
            "Start with the problem: manual defect detection is costly, and code vulnerabilities are subtle.",
            "Frame the research question: are prompt-only LLMs enough, or do we need task adaptation/code-specific models?",
            "Introduce the experimental ladder: zero-shot -> 4-shot -> LoRA fine-tuning -> TF-IDF baselines -> GraphCodeBERT.",
            "Show validation/test metrics and confusion matrices.",
            "Interpret the result: GraphCodeBERT has the best balanced Macro-F1; TF-IDF Logistic Regression is surprisingly strong; 4-shot collapse is a useful negative result.",
            "Close with limitations, future work, and reproducible repo/output locations.",
        ],
    )

    add_heading(doc, "4. Experimental Design", 1)
    add_table(
        doc,
        ["Model/Setting", "Purpose", "Adaptation Level", "Output Status"],
        [
            ["Qwen zero-shot", "Test raw LLM reasoning from prompt only", "None", "Validation + test complete"],
            ["Qwen 4-shot", "Test in-context learning with examples", "Prompt examples only", "Validation + test complete; collapsed to non-defective"],
            ["Qwen LoRA", "Task-adapt Qwen with efficient fine-tuning", "Supervised LoRA adapter", "Training + validation + test complete"],
            ["Majority baseline", "Sanity-check majority-class behavior", "Most frequent class", "Test complete"],
            ["TF-IDF Linear SVM", "Compare against a traditional lexical ML baseline", "Supervised CPU classifier", "Test complete"],
            ["TF-IDF Logistic Regression", "Compare against a strong lexical ML baseline", "Supervised CPU classifier", "Test complete"],
            ["GraphCodeBERT", "Compare against code-specific encoder", "Supervised encoder classifier", "Training + validation + test complete"],
        ],
        [1.45, 2.25, 1.45, 2.25],
    )

    add_heading(doc, "5. Evaluation Measures", 1)
    doc.add_paragraph(
        "Do not rely only on accuracy because defect detection is class-sensitive and the defective class is the most important target. "
        "Use Macro-F1 as the main balanced metric and Defective-F1/Recall to evaluate whether the model actually catches defective code."
    )
    add_bullets(
        doc,
        [
            "Accuracy: overall correctness, useful but can hide majority-class collapse.",
            "Macro Precision/Recall/F1: treats defective and non-defective classes more evenly.",
            "Defective Precision/Recall/F1: focuses on the positive class, which matters for security and code review.",
            "Confusion matrix: shows false positives and false negatives directly.",
        ],
    )

    add_heading(doc, "6. Final Results to Put in PPT", 1)
    add_table(
        doc,
        ["Experiment", "Split", "Accuracy", "Macro-F1", "Defective-F1", "Interpretation"],
        [
            ["Qwen zero-shot", "Validation", "0.5201", "0.5086", "0.4332", "Baseline LLM has limited signal"],
            ["Qwen zero-shot", "Test", "0.5264", "0.5180", "0.4545", "Prompt-only performance is modest"],
            ["Qwen 4-shot", "Validation", "0.5655", "0.3612", "0.0000", "Collapsed to non-defective"],
            ["Qwen 4-shot", "Test", "0.5406", "0.3509", "0.0000", "Useful negative result"],
            ["Qwen LoRA", "Validation", "0.5556", "0.5509", "0.5045", "Stable but limited improvement"],
            ["Qwen LoRA", "Test", "0.5523", "0.5507", "0.5236", "Main fine-tuned LLM result"],
            ["Majority baseline", "Test", "0.5406", "0.3509", "0.0000", "Sanity check; same collapse as 4-shot"],
            ["TF-IDF Linear SVM", "Test", "0.6007", "0.5994", "0.5766", "Strong lexical CPU baseline"],
            ["TF-IDF Logistic Regression", "Test", "0.6223", "0.6218", "0.6088", "Strongest traditional CPU baseline"],
            ["GraphCodeBERT", "Validation", "0.6618", "0.6513", "0.5908", "Best validation result"],
            ["GraphCodeBERT", "Test", "0.6589", "0.6517", "0.6017", "Best test result"],
        ],
        [1.45, 0.8, 0.8, 0.8, 0.9, 2.0],
    )

    add_note(
        doc,
        "Main conclusion",
        "For code defect detection, supervised baselines outperform prompt-only Qwen. TF-IDF Logistic Regression is surprisingly competitive, while GraphCodeBERT gives the strongest balanced test Macro-F1. The result suggests task supervision and code-aware modeling matter more than simply adding few-shot examples.",
    )

    add_heading(doc, "7. Suggested 10-Minute Demo Deck", 1)
    add_table(
        doc,
        ["Slide", "Title", "Time", "Speaker Goal"],
        [
            ["1", "Project Title + Team", "0:30", "Introduce team, task, and GitHub link."],
            ["2", "Problem & Motivation", "1:00", "Explain why defect detection matters and why this is hard."],
            ["3", "Research Question", "0:45", "Prompt-only LLM vs fine-tuned LLM vs code encoder."],
            ["4", "Dataset & Task", "1:00", "CodeXGLUE, binary labels, validation/test splits."],
            ["5", "Model Ladder", "1:15", "Zero-shot, 4-shot, LoRA, TF-IDF baselines, GraphCodeBERT."],
            ["6", "Evaluation Metrics", "0:45", "Macro-F1, Defective-F1, confusion matrix."],
            ["7", "Result Table", "1:30", "GraphCodeBERT has best Macro-F1; TF-IDF Logistic Regression is competitive; 4-shot collapse explained."],
            ["8", "Confusion Matrices", "1:00", "Show false positives/false negatives visually."],
            ["9", "Live Demo / Repo Walkthrough", "1:00", "Show GitHub repo, outputs, metrics files, scripts."],
            ["10", "Conclusion & Next Steps", "0:30", "State takeaway and future improvements."],
        ],
        [0.45, 2.05, 0.65, 3.0],
    )

    add_heading(doc, "8. GPT Prompt to Generate the PPT", 1)
    add_prompt_box(
        doc,
        """
Create a professional 10-minute PowerPoint deck for a university final project demo.

Project: Group 8 Code Defect Detection with LLMs and GraphCodeBERT.
Audience: instructor and classmates.
Tone: concise, technical, polished, demo-ready.
Team names: [fill names in alphabetical order by last name].
Code link: https://github.com/JiufuZh/526

Required slide flow:
1. Title + team + GitHub link.
2. Problem and motivation: detecting defective C/C++ functions is important for software quality and security.
3. Research question: are prompt-only LLMs enough, or do task adaptation and code-specific pretraining matter?
4. Dataset/task: CodeXGLUE defect detection; binary labels non-defective vs defective.
5. Methods: Qwen zero-shot, Qwen 4-shot, Qwen LoRA fine-tuning, majority baseline, TF-IDF Linear SVM, TF-IDF Logistic Regression, and GraphCodeBERT supervised encoder.
6. Evaluation: accuracy, macro precision/recall/F1, defective precision/recall/F1, confusion matrix.
7. Main results table using these numbers:
   - Zero-shot test: Accuracy 0.5264, Macro-F1 0.5180, Defective-F1 0.4545.
   - 4-shot test: Accuracy 0.5406, Macro-F1 0.3509, Defective-F1 0.0000.
   - LoRA test: Accuracy 0.5523, Macro-F1 0.5507, Defective-F1 0.5236.
   - Majority baseline test: Accuracy 0.5406, Macro-F1 0.3509, Defective-F1 0.0000.
   - TF-IDF Linear SVM test: Accuracy 0.6007, Macro-F1 0.5994, Defective-F1 0.5766.
   - TF-IDF Logistic Regression test: Accuracy 0.6223, Macro-F1 0.6218, Defective-F1 0.6088.
   - GraphCodeBERT test: Accuracy 0.6589, Macro-F1 0.6517, Defective-F1 0.6017.
8. Confusion matrix slide: emphasize that 4-shot and majority collapsed to non-defective, TF-IDF Logistic Regression catches many defective samples, and GraphCodeBERT has the best balanced Macro-F1.
9. Demo slide: show repo, scripts, metrics JSON files, figures, and reproducibility.
10. Conclusion: prompt-only LLMs are weak for this task; TF-IDF baselines are strong; GraphCodeBERT is the best balanced model.

Design requirements:
- Keep the deck to about 10 slides.
- Use clear charts/tables, not dense paragraphs.
- Make the result slide visually central.
- Include proper citations/URLs in small text on relevant slides.
- Avoid claiming perfect performance; emphasize limitations and future work.
- Include a backup slide with exact metrics and output file paths if useful.
""",
    )

    add_heading(doc, "9. Assets and File Paths", 1)
    add_bullets(
        doc,
        [
            "GitHub repo: https://github.com/JiufuZh/526",
            "Main result table: results/report_ready_metrics.md and results/report_ready_metrics.csv",
            "Confusion matrix overview: results/figures/confusion_matrices_overview.png",
            "GraphCodeBERT metrics: results/graphcodebert_validation_metrics.json and results/graphcodebert_test_metrics.json",
            "LoRA metrics: results/lora_validation_metrics.json and results/lora_test_metrics.json",
            "Zero-shot and 4-shot metrics: results/zero_shot_*_metrics.json and results/four_shot_*_metrics.json",
            "CPU baseline test metrics: results/majority_test_metrics.json, results/tfidf_linear_svm_test_metrics.json, and results/tfidf_logreg_test_metrics.json",
            "Tillicum project path: /gpfs/projects/imt526a/group8/final",
        ],
    )

    if FIG.exists():
        add_heading(doc, "10. Figure to Use in PPT", 1)
        doc.add_paragraph("Use this overview or the individual confusion matrix PNGs as high-resolution figures.")
        doc.add_picture(str(FIG), width=Inches(5.8))

    add_heading(doc, "11. References to Cite", 1)
    add_bullets(
        doc,
        [
            "CodeXGLUE benchmark: https://github.com/microsoft/CodeXGLUE",
            "GraphCodeBERT paper: https://arxiv.org/abs/2009.08366",
            "GraphCodeBERT model card: https://huggingface.co/microsoft/graphcodebert-base",
            "Qwen model family: https://huggingface.co/Qwen",
            "LoRA paper: https://arxiv.org/abs/2106.09685",
            "Project code and generated artifacts: https://github.com/JiufuZh/526",
        ],
    )

    add_heading(doc, "12. Report Checklist", 1)
    add_bullets(
        doc,
        [
            "No more than 8 pages; submit a single PDF for grading.",
            "Include team member names clearly, typically alphabetical by last name.",
            "Include citations/URLs for datasets, models, papers, code, and generated assets.",
            "Use sharp figures; do not use blurry screenshots.",
            "Do not overstate results: GraphCodeBERT is best among current experiments, not a solved defect detector.",
            "Only one group submission should be made.",
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Group 8 defect detection project - PPT generation brief")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
